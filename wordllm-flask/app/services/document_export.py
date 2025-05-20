"""
文档导出服务模块
提供将文档导出为不同格式的功能
"""
import os
import tempfile
from docx import Document as DocxDocument
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import logging

logger = logging.getLogger(__name__)

class DocumentExportService:
    """文档导出服务类，负责将文档内容转换为各种格式"""
    
    @staticmethod
    def get_number_style(chapter, style_config):
        """根据配置获取章节编号样式"""
        number_style = style_config.get('number_style', 'chapter')
        chapter_number = chapter.chapter_number
        
        if number_style == 'none':
            return ''
        elif number_style == 'chapter':
            # 使用中文样式
            if '.' not in str(chapter_number):
                return f"第{chapter_number}章 "
            elif str(chapter_number).count('.') == 1:
                return f"第{chapter_number}节 "
            else:
                # 一级编号使用汉字
                return f"{'一二三四五六七八九十'[int(chapter_number.split('.')[-1])-1]}、 "
        else:  # number style
            return f"{chapter_number} "
    
    @staticmethod
    def export_as_docx(project, chapters, settings):
        """
        将项目章节导出为DOCX格式
        
        Args:
            project: 项目模型实例
            chapters: 章节模型实例列表
            settings: 格式设置字典
        
        Returns:
            str: 临时文件路径
        """
        doc = DocxDocument()
        
        # 应用文档设置（页边距）
        section = doc.sections[0]
        margins = settings.get('margins', {})
        section.top_margin = Cm(float(margins.get('top', 2.54)))
        section.bottom_margin = Cm(float(margins.get('bottom', 2.54)))
        section.left_margin = Cm(float(margins.get('left', 3.18)))
        section.right_margin = Cm(float(margins.get('right', 3.18)))
        
        # 获取样式设置
        section_number_style = settings.get('section_number_style', {})
        title_styles = {
            1: settings.get('level1_style', {}),
            2: settings.get('level2_style', {}),
            3: settings.get('level3_style', {})
        }
        text_style = settings.get('text_style', {})
        
        # 添加标题和内容
        for chapter in chapters:
            # 确定标题级别
            level = 1
            if '.' in str(chapter.chapter_number):
                dot_count = str(chapter.chapter_number).count('.')
                level = dot_count + 1
                level = min(level, 3)  # 最多支持三级标题
            
            # 获取标题样式设置
            title_style = title_styles.get(level, {})
            
            # 添加编号和标题
            number_prefix = DocumentExportService.get_number_style(chapter, section_number_style)
            heading = doc.add_heading(level=level)
            heading_text = heading.add_run(f"{number_prefix}{chapter.title}")
            
            # 应用字体设置
            font_name = title_style.get('fontFamily', '仅宋体')
            if font_name == '仅宋体':
                font_name = '宋体'  # python-docx使用的名称
            heading_text.font.name = font_name
            
            # 应用字号设置
            font_size_mapping = {
                '小三': 15,
                '三号': 16,
                '小四': 12,
                '四号': 14,
                '小五': 9,
                '五号': 10.5
            }
            font_size = title_style.get('fontSize', '四号')
            heading_text.font.size = Pt(font_size_mapping.get(font_size, 14))
            
            # 应用加粗设置
            if title_style.get('bold', True):
                heading_text.font.bold = True
            
            # 应用对齐设置
            alignment_mapping = {
                '左对齐': WD_ALIGN_PARAGRAPH.LEFT,
                '居中': WD_ALIGN_PARAGRAPH.CENTER,
                '右对齐': WD_ALIGN_PARAGRAPH.RIGHT
            }
            alignment = title_style.get('alignment', '左对齐')
            heading.alignment = alignment_mapping.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
            
            # 添加内容
            if chapter.content:
                paragraphs = chapter.content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph()
                        
                        # 应用首行缩进
                        first_line_indent = text_style.get('firstLineIndent', 2)
                        para.paragraph_format.first_line_indent = Pt(first_line_indent * 10)
                        
                        # 应用对齐方式
                        text_alignment = text_style.get('alignment', '左对齐')
                        para.alignment = alignment_mapping.get(text_alignment, WD_ALIGN_PARAGRAPH.LEFT)
                        
                        # 添加文本
                        run = para.add_run(para_text)
                        
                        # 强制设置正文为宋体
                        run.font.name = '宋体'
                        # 兼容中英文
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        
                        # 应用字号
                        text_font_size = text_style.get('fontSize', '小四')
                        run.font.size = Pt(font_size_mapping.get(text_font_size, 12))
                        
                        # 应用加粗
                        if text_style.get('bold', False):
                            run.font.bold = True
        
        # 保存到临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        logger.info(f"文档已导出为DOCX格式: {temp_file.name}")
        return temp_file.name
    
    @staticmethod
    def export_as_txt(project, chapters, settings):
        """
        将项目章节导出为TXT格式
        
        Args:
            project: 项目模型实例
            chapters: 章节模型实例列表
            settings: 格式设置字典
        
        Returns:
            str: 临时文件路径
        """
        # 获取章节编号样式
        section_number_style = settings.get('section_number_style', {})
        
        # 创建临时文件
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.txt', mode='w', encoding='utf-8')
        
        # 写入内容
        for chapter in chapters:
            # 添加编号和标题
            number_prefix = DocumentExportService.get_number_style(chapter, section_number_style)
            temp_file.write(f"{number_prefix}{chapter.title}\n\n")
            
            # 添加内容
            if chapter.content:
                temp_file.write(f"{chapter.content}\n\n")
        
        temp_file.close()
        logger.info(f"文档已导出为TXT格式: {temp_file.name}")
        return temp_file.name
    
    @staticmethod
    def convert_to_pdf(docx_path):
        """
        将DOCX转换为PDF (这需要额外的依赖项如libreoffice或者其他PDF转换库)
        此处为简化实现，实际项目中可能需要更复杂的转换逻辑
        
        Args:
            docx_path: DOCX文件路径
        
        Returns:
            str: PDF临时文件路径
        """
        # 注意：这里仅作为示例，实际项目中需要实现具体的转换逻辑
        # 你可以使用如下方法：
        # 1. 使用libreoffice命令行工具
        # 2. 使用docx2pdf库
        # 3. 使用python-docx-replace和ReportLab结合的方式
        
        # 创建具有相同文件名但扩展名为.pdf的临时文件
        pdf_path = docx_path.replace('.docx', '.pdf')
        
        # 这里应该有实际的转换代码
        logger.warning("PDF转换功能尚未完全实现，需要额外依赖")
        
        # 由于没有实际转换，我们返回原始DOCX路径
        # 在实际项目中，这里应该返回转换后的PDF路径
        return docx_path
